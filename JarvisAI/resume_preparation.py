"""
Module for resume preparation using templates.
"""

from docx import Document

class ResumePreparation:
    def __init__(self):
        self.document = Document()

    def create_resume(self, name, email, phone, experiences, education):
        self.document.add_heading(name, 0)
        self.document.add_paragraph(f"Email: {email}")
        self.document.add_paragraph(f"Phone: {phone}")

        self.document.add_heading('Experience', level=1)
        for exp in experiences:
            self.document.add_paragraph(exp, style='ListBullet')

        self.document.add_heading('Education', level=1)
        for edu in education:
            self.document.add_paragraph(edu, style='ListBullet')

    def save_resume(self, filename):
        self.document.save(filename)
